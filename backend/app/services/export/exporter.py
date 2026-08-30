import io
import csv
from typing import List, Dict, Any, Optional
from datetime import timedelta


class TranscriptExporter:
    """Generates structured exports for transcripts: TXT, CSV, SRT, PDF, DOCX."""

    @staticmethod
    def format_timestamp_short(seconds: float) -> str:
        """Format seconds to MM:SS or HH:MM:SS."""
        total_seconds = int(seconds)
        hours = total_seconds // 3600
        mins = (total_seconds % 3600) // 60
        secs = total_seconds % 60
        if hours > 0:
            return f"{hours:02d}:{mins:02d}:{secs:02d}"
        return f"{mins:02d}:{secs:02d}"

    @staticmethod
    def format_timestamp_srt(seconds: float) -> str:
        """Format seconds to SRT standard format: HH:MM:SS,mmm"""
        hrs = int(seconds // 3600)
        mins = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int(round((seconds - int(seconds)) * 1000))
        if millis >= 1000:
            secs += 1
            millis = 0
        return f"{hrs:02d}:{mins:02d}:{secs:02d},{millis:03d}"

    @classmethod
    def export_txt(
        cls,
        title: str,
        segments: List[Dict[str, Any]],
        include_timestamps: bool = True
    ) -> str:
        """Generates plain text transcript."""
        lines = [f"# {title}", ""]
        for seg in segments:
            text = seg.get("text", "").strip()
            speaker = seg.get("speaker")
            start = seg.get("start_time", 0.0)
            time_str = cls.format_timestamp_short(start)

            if include_timestamps:
                speaker_str = f" [{speaker}]" if speaker else ""
                lines.append(f"{time_str}{speaker_str}")
                lines.append(text)
                lines.append("")
            else:
                speaker_prefix = f"[{speaker}]: " if speaker else ""
                lines.append(f"{speaker_prefix}{text}")

        return "\n".join(lines).strip()

    @classmethod
    def export_csv(cls, segments: List[Dict[str, Any]]) -> str:
        """Generates CSV transcript with columns: Timestamp, Speaker, Transcript."""
        output = io.StringIO()
        writer = csv.writer(output, lineterminator="\n")
        writer.writerow(["Timestamp", "Speaker", "Transcript"])

        for seg in segments:
            time_str = cls.format_timestamp_short(seg.get("start_time", 0.0))
            speaker = seg.get("speaker") or ""
            text = seg.get("text", "").strip()
            writer.writerow([time_str, speaker, text])

        return output.getvalue()

    @classmethod
    def export_srt(cls, segments: List[Dict[str, Any]]) -> str:
        """Generates standard SubRip (.srt) subtitle file content."""
        lines = []
        for idx, seg in enumerate(segments, 1):
            start = cls.format_timestamp_srt(seg.get("start_time", 0.0))
            end = cls.format_timestamp_srt(seg.get("end_time", seg.get("start_time", 0.0) + 2.0))
            text = seg.get("text", "").strip()
            speaker = seg.get("speaker")
            if speaker:
                text = f"[{speaker}] {text}"

            lines.append(str(idx))
            lines.append(f"{start} --> {end}")
            lines.append(text)
            lines.append("")

        return "\n".join(lines).strip()

    @classmethod
    def export_pdf(
        cls,
        title: str,
        segments: List[Dict[str, Any]],
        summary: Optional[str] = None
    ) -> bytes:
        """Generates formatted PDF document."""
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=20,
            leading=24,
            textColor=colors.HexColor('#0f172a'),
            spaceAfter=12
        )
        meta_style = ParagraphStyle(
            'MetaStyle',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#64748b'),
            spaceAfter=16
        )
        time_style = ParagraphStyle(
            'TimeStyle',
            parent=styles['Normal'],
            fontSize=9,
            leading=12,
            textColor=colors.HexColor('#2563eb'),
            fontName='Helvetica-Bold'
        )
        text_style = ParagraphStyle(
            'TextStyle',
            parent=styles['Normal'],
            fontSize=10,
            leading=15,
            textColor=colors.HexColor('#1e293b'),
            spaceAfter=8
        )
        summary_style = ParagraphStyle(
            'SummaryStyle',
            parent=styles['Normal'],
            fontSize=10,
            leading=15,
            textColor=colors.HexColor('#334155'),
            spaceAfter=12
        )

        story = []
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"AI Video & Voice Transcriber — Transcript Document", meta_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0'), spaceAfter=14))

        if summary:
            story.append(Paragraph("<b>Executive Summary</b>", styles['Heading3']))
            story.append(Paragraph(summary, summary_style))
            story.append(Spacer(1, 10))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#cbd5e1'), spaceAfter=14))

        story.append(Paragraph("<b>Full Transcript</b>", styles['Heading3']))
        story.append(Spacer(1, 6))

        for seg in segments:
            start_str = cls.format_timestamp_short(seg.get("start_time", 0.0))
            speaker = seg.get("speaker")
            speaker_tag = f" &bull; {speaker}" if speaker else ""
            time_line = f"<b>{start_str}</b>{speaker_tag}"
            text = seg.get("text", "")

            story.append(Paragraph(time_line, time_style))
            story.append(Paragraph(text, text_style))

        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    @classmethod
    def export_docx(
        cls,
        title: str,
        segments: List[Dict[str, Any]],
        summary: Optional[str] = None
    ) -> bytes:
        """Generates Microsoft Word (.docx) document."""
        import docx
        from docx.shared import Pt, RGBColor

        doc = docx.Document()
        
        # Heading
        h1 = doc.add_heading(title, level=1)
        
        # Subtitle
        p = doc.add_paragraph()
        p_run = p.add_run("AI Video & Voice Transcriber — Verified Transcript")
        p_run.font.color.rgb = RGBColor(100, 116, 139)
        p_run.font.size = Pt(10)

        # Summary if available
        if summary:
            doc.add_heading("Executive Summary", level=2)
            doc.add_paragraph(summary)

        doc.add_heading("Full Transcript", level=2)

        for seg in segments:
            time_str = cls.format_timestamp_short(seg.get("start_time", 0.0))
            speaker = seg.get("speaker")
            speaker_label = f" [{speaker}]" if speaker else ""
            
            p_time = doc.add_paragraph()
            run_time = p_time.add_run(f"{time_str}{speaker_label}")
            run_time.bold = True
            run_time.font.color.rgb = RGBColor(37, 99, 235)
            
            p_text = doc.add_paragraph(seg.get("text", ""))
            p_text.paragraph_format.space_after = Pt(8)

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()
